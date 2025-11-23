import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import datetime
import re

def read_sheets(path):
    xls = pd.ExcelFile(path)
    sheets = {name: xls.parse(name) for name in xls.sheet_names}
    return sheets

def set_cell_border(cell, **kwargs):
    # helper omitted for brevity (python-docx table border helper).
    pass

def _add_table_borders(table):
    # apply borders to the table
    tbl = table._tbl
    tblPr = tbl.get_or_add_tblPr()
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top','left','bottom','right','insideH','insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # width
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def write_header(document, title_text, client_name):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Incremint Edge Pvt Ltd\n")
    run.bold = True
    run.font.size = Pt(14)
    run = p.add_run(title_text + "\n")
    run.bold = True
    run.font.size = Pt(16)
    if client_name:
        p = document.add_paragraph()
        p.add_run(f"Client: {client_name}\n")
    p = document.add_paragraph()
    p.add_run(f"Date: {datetime.datetime.now().strftime('%d-%m-%Y %H:%M')}\n")

def add_client_details_table(document, client_df):
    document.add_heading("Client Details", level=3)
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Field"
    hdr_cells[1].text = "Value"
    for col in client_df.columns:
        val = str(client_df.iloc[0].get(col, ""))
        row = table.add_row().cells
        row[0].text = col
        row[1].text = val
    _add_table_borders(table)

def add_cover_details_table(document, cover_info):
    document.add_heading("Cover Details", level=3)
    table = document.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Sum Assured"
    hdr[1].text = "Policy Term"
    hdr[2].text = "Cover Till Age"
    hdr[3].text = "PPT"
    # single row fill
    r = table.add_row().cells
    r[0].text = str(cover_info.get("Sum Assured", ""))
    r[1].text = str(cover_info.get("Policy Term", ""))
    r[2].text = str(cover_info.get("Cover Till Age", ""))
    r[3].text = str(cover_info.get("PPT", ""))
    _add_table_borders(table)

def add_premium_comparison(document, premiums_df):
    document.add_heading("Premium Comparison (Regular/10 Pay)", level=3)
    # wide first column for logos as placeholder
    cols = [" ", "Company", "Plan", "Regular Premium", "10 Pay Premium", "Notes"]
    table = document.add_table(rows=1, cols=len(cols))
    hdr_cells = table.rows[0].cells
    for i, name in enumerate(cols):
        hdr_cells[i].text = name
    for _, row in premiums_df.iterrows():
        r = table.add_row().cells
        r[0].text = ""  # logo placeholder (wide col)
        r[1].text = str(row.get("Insurance Company", ""))
        r[2].text = str(row.get("Plan Name", ""))
        r[3].text = str(row.get("Regular Premium", ""))
        r[4].text = str(row.get("10 Pay Premium", row.get("10 Pay", "")))
        r[5].text = str(row.get("Special Notes", ""))
    _add_table_borders(table)

def add_advisory_note(document, notes_text):
    document.add_heading("Advisory Note", level=3)
    p = document.add_paragraph()
    p.add_run(notes_text)

def make_term_quote_from_excel(excel_path):
    # Read sheets
    sheets = read_sheets(excel_path)
    # Default sheet handling
    client_df = sheets.get("Client Details", None)
    premiums_df = sheets.get("Premiums", None)
    final_notes_df = sheets.get("Final Notes", None)

    # Create document
    doc = Document()
    client_name = ""
    if client_df is not None and not client_df.empty:
        # try detect client name column
        if "Client Name" in client_df.columns:
            client_name = str(client_df.loc[0, "Client Name"])
    write_header(doc, "Final Incremint Dual-Pay Term Quote", client_name)

    # Client table
    if client_df is not None and not client_df.empty:
        add_client_details_table(doc, client_df)

    # Cover details (try to infer)
    cover_info = {}
    if client_df is not None and not client_df.empty:
        # try common columns
        for key in ["Sum Assured", "Policy Term", "Cover Till Age", "PPT"]:
            if key in client_df.columns:
                cover_info[key] = client_df.loc[0, key]
    add_cover_details_table(doc, cover_info)

    # Premiums
    if premiums_df is not None:
        # Normalize column names if needed
        # Expect columns: Insurance Company, Plan Name, Regular Premium, 10 Pay Premium, Special Notes
        add_premium_comparison(doc, premiums_df)

    # Advisory
    notes_text = ""
    if final_notes_df is not None and not final_notes_df.empty:
        # combine all text cells
        notes_text = " ".join(map(str, final_notes_df.fillna("").astype(str).apply(lambda row: " ".join(row), axis=1).tolist()))
    else:
        notes_text = "Recommendation: 10 Pay offers quicker benefit accumulation — consider max allowed cover and fixed premiums for life. Contact your advisor for exact tailored recommendation. ⚠️ This is a system-generated quote."
    add_advisory_note(doc, notes_text)

    # Contact section
    doc.add_heading("Contact Now 📞", level=3)
    p = doc.add_paragraph()
    p.add_run("Agent Name: __________________\nMobile: __________________\n")

    # Save file
    out_fname = os.path.splitext(os.path.basename(excel_path))[0] + "_term_quote_final.docx"
    out_path = os.path.join("/tmp", out_fname)
    doc.save(out_path)
    return out_path

if __name__ == "__main__":
    # quick local test placeholder
    print("Module loaded")
